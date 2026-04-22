"""
Newfoundland and Labrador Powerlifting Association provincial QT scraper.

NLPA publishes qualifying totals as a Google-hosted .docx. Each of the 8
tables in the document covers one (sex, equipment, event) combination,
preceded by a paragraph label (e.g. "Men's Classic 3 Lift"). The
scraper walks document body order to match tables with their labels.

As of 2026-04, the latest file still has a 2022 creation date and title
("Provincial Qualifying Totals" with no year). Treat the data as
effective-year=2022 and log a staleness warning when the source doc is
more than 24 months old so a future refresh is visible at scrape time.
"""
from __future__ import annotations

import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

import docx
import requests
from docx.oxml.ns import qn

log = logging.getLogger(__name__)

NLPA_LANDING_URL = "https://www.nlpowerlifting.ca/athletes/qualifying-totals"
NLPA_DOC_ID = "1064rndgGmi9X_Ebf8RgH5dg7OljNmGgQ"
_USER_AGENT = "cpu-analytics-scraper/1.0 (+contact matthias.bernhard7@gmail.com)"
_HTTP_TIMEOUT = 30.0

# Staleness threshold in days. When the source doc's creation date is
# older than this the scraper logs a WARNING so the orchestrator's log
# surfaces a refresh reminder without failing the run.
STALENESS_WARN_DAYS = 730  # 2 years

# Column order matches the 7 columns following the weight-class cell.
_DIVISION_COLUMNS = (
    "Open", "Sub-Junior", "Junior",
    "Master 1", "Master 2", "Master 3", "Master 4",
)


def _export_url(doc_id: str = NLPA_DOC_ID) -> str:
    return f"https://docs.google.com/document/d/{doc_id}/export?format=docx"


def download_docx(target_dir: Path, doc_id: str = NLPA_DOC_ID) -> Path:
    """Download the NLPA docx via Google Docs export."""
    target_dir.mkdir(parents=True, exist_ok=True)
    dst = target_dir / "nlpa_qualifying_totals.docx"
    r = requests.get(
        _export_url(doc_id),
        headers={"User-Agent": _USER_AGENT},
        timeout=_HTTP_TIMEOUT,
        allow_redirects=True,
    )
    r.raise_for_status()
    dst.write_bytes(r.content)
    log.info("downloaded nlpa docx (%d bytes)", dst.stat().st_size)
    return dst


def _iter_body(doc: docx.Document) -> Iterator[tuple[str, object]]:
    """Yield (kind, element) pairs in document body order.

    ``kind`` is "paragraph" or "table"; the element is the matching
    docx.Paragraph / docx.Table wrapper.
    """
    p_iter = iter(doc.paragraphs)
    t_iter = iter(doc.tables)
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("paragraph", next(p_iter))
        elif child.tag == qn("w:tbl"):
            yield ("table", next(t_iter))


def _label_to_equipment_event(label: str) -> tuple[str, str] | None:
    """Parse "Men's Classic 3 Lift" / "Women's Equipped Bench Only" etc.

    Returns (equipment, event) or None for unrelated paragraphs.
    """
    s = label.lower().replace("\u2019", "'")  # curly quote -> straight
    equipment: str | None = None
    if "classic" in s:
        equipment = "Classic"
    elif "equipped" in s:
        equipment = "Equipped"
    if equipment is None:
        return None
    if "bench only" in s:
        return (equipment, "B")
    if "3 lift" in s or "3-lift" in s:
        return (equipment, "SBD")
    return None


def _normalise_weight_class(raw: str) -> str | None:
    s = (raw or "").strip().replace("+", "+")
    if not s:
        return None
    if not re.match(r"^\d+\+?$", s):
        return None
    return s


def _parse_qt(cell_text: str) -> float | None:
    """Parse a QT cell. Handle blanks, dashes, and stray commas from the
    docx authoring ("32,5" used instead of "32.5")."""
    s = (cell_text or "").strip().replace("\xa0", "")
    if s in ("", "-", "—", "–"):
        return None
    # Comma-as-decimal (one typo observed in the 2022 doc).
    s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _effective_year_from_doc(doc: docx.Document, fallback: int) -> int:
    """Pull the effective year from the title or fall back to creation
    year. The 2022 file has no year in the title text."""
    for p in doc.paragraphs:
        m = re.search(r"(20\d{2})\s+Provincial\s+Qualifying", p.text)
        if m is not None:
            return int(m.group(1))
    return fallback


def _creation_year(doc: docx.Document, fallback: int) -> int:
    """The file creation timestamp. Used as the effective year when the
    document title omits it (which is the 2022 NLPA state)."""
    created = doc.core_properties.created
    if created is not None:
        return created.year
    return fallback


def _is_stale(doc: docx.Document) -> bool:
    created = doc.core_properties.created
    if created is None:
        return False
    now = datetime.now(tz=timezone.utc)
    age_days = (now - created).days
    return age_days > STALENESS_WARN_DAYS


def parse_docx(docx_path: Path) -> list[dict]:
    """Parse the NLPA docx into QT rows (Classic + Equipped x SBD + B).

    The orchestrator filters out Equipped and Bench Only downstream.
    """
    doc = docx.Document(docx_path)
    effective_year = _effective_year_from_doc(
        doc, fallback=_creation_year(doc, fallback=2022),
    )
    if _is_stale(doc):
        log.warning(
            "NLPA docx is stale (created %s); provincial totals may be "
            "out of date",
            doc.core_properties.created,
        )

    rows: list[dict] = []
    current_equipment: str | None = None
    current_event: str | None = None
    for kind, element in _iter_body(doc):
        if kind == "paragraph":
            label = element.text
            mapping = _label_to_equipment_event(label)
            if mapping is not None:
                current_equipment, current_event = mapping
            continue

        # kind == "table"
        if current_equipment is None or current_event is None:
            continue
        header_cells = [c.text.strip().lower() for c in element.rows[0].cells]
        if not header_cells:
            continue
        if header_cells[0].startswith("men"):
            sex = "M"
        elif header_cells[0].startswith("women"):
            sex = "F"
        else:
            continue

        for row in element.rows[1:]:
            cells = [c.text for c in row.cells]
            if not cells:
                continue
            wc = _normalise_weight_class(cells[0])
            if wc is None:
                continue
            for col_idx, division in enumerate(_DIVISION_COLUMNS, start=1):
                if col_idx >= len(cells):
                    break
                qt = _parse_qt(cells[col_idx])
                if qt is None:
                    continue
                rows.append({
                    "sex": sex,
                    "level": "Provincials",
                    "region": None,
                    "division": division,
                    "equipment": current_equipment,
                    "event": current_event,
                    "weight_class": wc,
                    "qt": qt,
                    "effective_year": effective_year,
                    "province": "Newfoundland and Labrador",
                })

    log.info("parsed %d rows from nlpa docx", len(rows))
    return rows
